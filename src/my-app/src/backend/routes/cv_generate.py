from flask import Blueprint, request, send_file
from docx import Document
from io import BytesIO
from flask_cors import cross_origin 

cv_generate_bp = Blueprint('cv_generate', __name__)

@cv_generate_bp.route('/generate-cv', methods=['POST', 'OPTIONS'])
@cross_origin()
def generate_cv():
    if request.method == 'OPTIONS':
        return '', 200

    data = request.json
    doc = Document()
    
    doc.add_heading(data.get("name", "No Name"), 0)
    
    if data.get("location"):
        doc.add_paragraph(data["location"])

    if data.get("contact"):
        doc.add_paragraph(data["contact"])
    
    def add_section(title, content):
        if content is not None and str(content).strip() != "":
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)
     
    add_section("Education", data.get("education"))
    add_section("Experience", data.get("experience"))
    add_section("Projects", data.get("projects"))
    add_section("Leadership", data.get("leadership"))
    add_section("Skills", data.get("skills"))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name="Updated_CV.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
